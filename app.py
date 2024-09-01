import streamlit as st
import requests
import json
from io import BytesIO
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt

# Set page configuration
st.set_page_config(page_title="Análisis de Tendencias Actuales", page_icon="📊")

# Titles and Main Column
st.title("Análisis de Tendencias Actuales")

# Describe app functionality
st.write("""
    Esta aplicación busca información sobre tendencias actuales utilizando la API de Serper y genera un análisis o 
    informe sobre dichas tendencias utilizando la API de Together.
""")

if st.button("Buscar Tendencias y Generar Informe"):
    with st.spinner("Buscando tendencias actuales..."):
        # Fetch trends using Serper API
        serper_api_key = st.secrets["SERPER_API_KEY"]
        serper_url = "https://api.serper.dev/search"
        headers = {
            "Content-Type": "application/json",
            "X-API-KEY": serper_api_key
        }
        query = "current trends"
        payload = json.dumps({
            "q": query,
            "gl": "us",
        })
        response = requests.post(serper_url, headers=headers, data=payload)
        trends = response.json()

        trend_titles = []
        for item in trends.get('organic', []):
            trend_titles.append(item['title'])

        if not trend_titles:
            st.error("No se encontraron tendencias actuales.")
        else:
            st.success("Tendencias encontradas. Generando informe...")

            # Generate analysis report on trends using Together API
            together_api_key = st.secrets["TOGETHER_API_KEY"]
            together_url = "https://api.together.xyz/inference"
            trends_text = "\n".join(f"- {title}" for title in trend_titles)
            prompt = f"""
                Genera un informe detallado sobre las siguientes tendencias actuales:

                {trends_text}

                Proporciona un análisis profundo de cada tendencia, incluyendo posibles impactos sociales, económicos y 
                tecnológicos.
            """
            payload = json.dumps({
                "model": "mistralai/Mixtral-8x7B-Instruct-v0.1",
                "prompt": prompt,
                "max_tokens": 4096,
                "temperature": 0.7,
                "top_p": 0.9,
                "top_k": 50,
                "repetition_penalty": 1.1
            })
            headers = {
                'Authorization': f'Bearer {together_api_key}',
                'Content-Type': 'application/json'
            }
            response = requests.post(together_url, headers=headers, data=payload)
            analysis_text = response.json()['output']['choices'][0]['text'].strip()
            st.success("Informe generado con éxito.")
            st.write(analysis_text)

            # Create and download DOCX report
            if analysis_text:
                doc = Document()

                # Define styles
                styles = doc.styles
                if 'Sin Sangría' not in styles:
                    style = styles.add_style('Sin Sangría', WD_STYLE_TYPE.PARAGRAPH)
                    style.font.name = 'Calibri'
                    style.font.size = Pt(11)
                    style.paragraph_format.space_after = Pt(10)
                    style.paragraph_format.left_indent = Pt(0)

                doc.add_heading("Informe de Tendencias Actuales", 0)
                doc.add_heading("Tendencias Encontradas", level=1)

                for title in trend_titles:
                    doc.add_paragraph(f"- {title}", style='Sin Sangría')

                doc.add_heading("Análisis de Tendencias", level=1)
                doc.add_paragraph(analysis_text, style='Sin Sangría')

                doc.add_paragraph('\nNota: Este informe fue generado por un asistente de IA. Se recomienda revisar y editar el contenido para garantizar su precisión y calidad.', style='Sin Sangría')

                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                st.download_button(
                    label="Descargar informe en DOCX",
                    data=buffer,
                    file_name="informe_tendencias_actuales.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
